import io
from dataclasses import dataclass, field

import pdfplumber
from docx import Document

from cv_parser.leaves.extraction_leaves import page_area
from cv_parser.models import ExtractionMethod


@dataclass
class RawExtraction:
    text: str
    method: ExtractionMethod
    chars: list[dict] = field(default_factory=list)
    total_page_area: float = 0.0


_SUPPORTED_EXTS = ("pdf", "docx", "txt")


# -------------------- extract_pdf ----------- START ----------
# -- Calls : page_area
# -- Called by: dispatch_by_extension
def extract_pdf(file_bytes: bytes) -> RawExtraction:
    pages_text: list[str] = []
    chars: list[dict] = []
    total_area = 0.0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            total_area += page_area(float(page.width), float(page.height))
            page_text = page.extract_text() or ""
            pages_text.append(page_text)
            for ch in page.chars:
                chars.append({
                    "text": ch.get("text", ""),
                    "x0": float(ch.get("x0", 0.0)),
                    "top": float(ch.get("top", 0.0)),
                    "page_number": page.page_number,
                })

    text = "\n".join(pages_text).strip()
    method: ExtractionMethod = "digital_pdf" if text else "ocr_pdf"
    return RawExtraction(text=text, method=method, chars=chars, total_page_area=total_area)
# -------------------- extract_pdf ------------- END ----------------


# -------------------- extract_docx ----------- START ----------
# -- Calls : nothing (leaf-ish: only docx lib)
# -- Called by: dispatch_by_extension
def extract_docx(file_bytes: bytes) -> RawExtraction:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return RawExtraction(text="\n".join(parts).strip(), method="docx")
# -------------------- extract_docx ------------- END ----------------


# -------------------- extract_text ----------- START ----------
# -- Calls : nothing
# -- Called by: dispatch_by_extension
def extract_text(file_bytes: bytes) -> RawExtraction:
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="replace")
    return RawExtraction(text=text.strip(), method="text")
# -------------------- extract_text ------------- END ----------------


# -------------------- _extension_of ----------- START ----------
# -- Calls : nothing (leaf)
# -- Called by: dispatch_by_extension
def _extension_of(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
# -------------------- _extension_of ------------- END ----------------


# -------------------- dispatch_by_extension ----------- START ----------
# -- Calls : _extension_of, extract_pdf, extract_docx, extract_text
# -- Called by: parse_cv (parser.py)
def dispatch_by_extension(file_bytes: bytes, filename: str) -> RawExtraction:
    ext = _extension_of(filename)
    if ext == "pdf":
        return extract_pdf(file_bytes)
    if ext == "docx":
        return extract_docx(file_bytes)
    if ext == "txt":
        return extract_text(file_bytes)
    return RawExtraction(text="", method="failed")
# -------------------- dispatch_by_extension ------------- END ----------------


# -------------------- supported_extensions ----------- START ----------
# -- Calls : nothing (leaf)
# -- Called by: parse_endpoint (interface.py)
def supported_extensions() -> tuple[str, ...]:
    return _SUPPORTED_EXTS
# -------------------- supported_extensions ------------- END ----------------
